from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/jonathanmuhire/Documents/Oppprep")
DOCX = OUT / "opportunity-intelligence-os-brief.docx"
GDOCS_DOCX = OUT / "opportunity-intelligence-os-brief.gdocs.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_inches):
    width = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_inches=6.5):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def style_run(run, size=None, bold=None, color="000000"):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(title)
    style_run(run, size=26, bold=False)

    s = doc.add_paragraph()
    s.paragraph_format.space_after = Pt(14)
    r = s.add_run(subtitle)
    style_run(r, size=11, color="555555")


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=20, bold=False)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=16, bold=False)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=14, bold=False, color="434343")


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2)
    else:
        r = p.add_run(text)
        style_run(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    style_run(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    style_run(r)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)

    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_width(hdr[i], widths[i])
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr[i], "FFFFFF")
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        style_run(run, size=10, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            style_run(run, size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.color.rgb = RGBColor(0, 0, 0)

    return doc


def build():
    doc = setup_doc()

    add_title(
        doc,
        "Opportunity Intelligence OS",
        "A multi-agent workflow for turning technical curiosity into targeted opportunity, credible artifacts, and human-approved outreach.",
    )

    add_h1(doc, "Rough Problem Description")
    add_para(
        doc,
        "Talented students and early-career technical people often know they want to work on meaningful problems, but the job-search process pushes them toward generic applications, shallow networking, and cold messages that do not prove technical taste. The missing layer is not more job boards. It is a repeatable system for discovering real team problems, understanding them deeply enough to contribute, identifying the right people, and producing a credible note before outreach.",
    )
    add_para(
        doc,
        "This project is an opportunity intelligence and outreach operating system. It should help a user move from broad interest areas to company-specific technical packets: who to contact, what problem to discuss, what artifact to show, how the user's adjacent experience maps to the problem, and when to send or follow up.",
    )

    add_h1(doc, "Product Thesis")
    add_para(
        doc,
        "The system should behave like a small research-and-outreach team. It should not behave like a spam sender. The user remains the final decision-maker, especially before external actions such as LinkedIn connection requests, emails, scheduling, or messages.",
    )
    for item in [
        "Primary user: a technical student or early-career builder trying to break into high-signal teams.",
        "Primary output: a company opportunity packet, not just a message draft.",
        "Success metric: high-quality technical conversations started, not message volume.",
        "Operating principle: action over confusion, but never at the cost of fabricated claims or uncontrolled outreach.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "Core Workflow")
    steps = [
        "Source target companies by lane: AI infrastructure, inference systems, agentic AI, developer tools, data platforms, robotics AI, and adjacent technical areas.",
        "Source open problems from company blogs, papers, docs, GitHub, job descriptions, LinkedIn posts, Hacker News, Reddit, conference talks, and public technical discussions.",
        "Source relevant people who are close to the problem: engineers, researchers, founders, FDEs, product engineers, hiring managers, and recruiters.",
        "Write a clear technical note explaining the problem, the current landscape, and one concrete contribution the user could build or reason about.",
        "Map adjacent proof from the user's background. This can include prior projects, GitHub repositories, research, systems work, evaluation work, product judgment, or a closely related problem-solving pattern.",
        "Choose the outreach channel: school email, LinkedIn, conference follow-up, X, or a warm intro when available.",
        "Draft outreach under 200 words. Write like a student: specific, plain, human, and respectful. Avoid inflated claims, hype, and weird phrasing.",
        "Run verification, ask for approval, send only the approved message, then track status and follow-up.",
    ]
    for step in steps:
        add_number(doc, step)

    add_h1(doc, "Target Packet")
    add_para(doc, "Each company should produce one packet that can be reviewed quickly and reused across outreach.")
    packet_rows = [
        ("Company fit", "Why this company belongs on the list; target lane; role fit; sponsorship or hiring notes when relevant."),
        ("Open problem", "One specific problem the team appears to care about, backed by source URLs."),
        ("People map", "Three to eight people ranked by relevance, with public profile links and why each person matters."),
        ("Technical note", "One page explaining the problem, landscape, possible contribution, and evaluation plan."),
        ("Adjacent proof", "A sober mapping from the user's background to the problem without overstating experience."),
        ("Outreach drafts", "LinkedIn note, email, accepted-connection follow-up, and optional recruiter/FDE variant."),
        ("Verification", "Source checks, word counts, claim checks, no guessed emails unless explicitly approved."),
        ("CRM status", "Prepared, approved, sent, pending, accepted, replied, follow-up due, archived."),
    ]
    add_table(doc, ["Packet section", "Purpose"], packet_rows, [1.65, 4.85])

    add_h1(doc, "Multi-Agent Architecture")
    add_para(
        doc,
        "The architecture should separate research, synthesis, writing, verification, and external action. The orchestrator owns state and decides which specialist agent should run next.",
    )
    agent_rows = [
        ("Orchestrator", "Owns the task graph, state, approvals, retries, budget, and handoffs. It decides what can be automated and what needs user approval."),
        ("Profile Agent", "Builds the user's technical map from resume, GitHub, projects, interests, constraints, and preferred roles."),
        ("Company Agent", "Sources companies and scores them by technical fit, hiring relevance, sponsor likelihood, and reachable people."),
        ("Problem Agent", "Extracts current problems from public sources and produces source-backed problem briefs."),
        ("People Agent", "Finds and ranks relevant people by proximity to the problem and likely usefulness of outreach."),
        ("Technical Note Agent", "Writes one-page technical notes with problem framing, landscape, contribution idea, and evaluation approach."),
        ("Outreach Agent", "Writes concise human messages for email, LinkedIn, conference follow-up, and accepted-connection follow-up."),
        ("QA Agent", "Checks sources, claims, word count, tone, duplicated outreach, missing approvals, and unsupported experience claims."),
        ("Action Agent", "Creates drafts, opens browser tasks, sends only after explicit approval, and records outcomes."),
    ]
    add_table(doc, ["Agent", "Responsibility"], agent_rows, [1.55, 4.95])

    add_h1(doc, "Prioritized Execution Plan")
    add_h2(doc, "P0: Manual Gold Standard")
    add_para(
        doc,
        "Create ten high-quality packets manually or semi-manually. Baseten is the reference example. This phase defines what good looks like before automating.",
    )
    for item in [
        "Deliverable: 10 company packets with sources, people, one-pagers, and outreach drafts.",
        "Gate: every packet must be strong enough that the user would actually send the first message.",
        "Avoid: building agent complexity before the packet format is proven.",
    ]:
        add_bullet(doc, item)

    add_h2(doc, "P1: Data Model and Tracker")
    add_para(
        doc,
        "Create a lightweight CRM and artifact store. Start with local files plus SQLite or a simple structured store before adding a larger backend.",
    )
    for item in [
        "Entities: user profile, company, problem, person, source, packet, message, approval, send event, follow-up.",
        "Statuses: sourced, drafted, verified, approved, sent, pending, accepted, replied, follow-up due, closed.",
        "Artifacts: one-pagers, source notes, outreach variants, verification logs.",
    ]:
        add_bullet(doc, item)

    add_h2(doc, "P2: Agent Harness")
    add_para(
        doc,
        "Implement the orchestrator as a task graph. Each agent receives a clear input object and returns structured JSON plus a human-readable artifact. The first agent harness does not need to be fancy. It needs to be observable and reliable.",
    )
    for item in [
        "Every agent output includes confidence, source URLs, assumptions, and next recommended action.",
        "The QA agent runs before any user-facing artifact is considered ready.",
        "The Action Agent is disabled by default for external sends and requires explicit approval.",
    ]:
        add_bullet(doc, item)

    add_h2(doc, "P3: Connectors")
    add_para(
        doc,
        "Add integrations only after the core packet loop works. The connectors are useful because they reduce friction, not because they replace judgment.",
    )
    connector_rows = [
        ("GitHub", "Read user's repositories and generate adjacent proof."),
        ("Browser", "Research company pages, LinkedIn profiles, blogs, and public discussions."),
        ("Google Docs/Drive", "Create polished one-pagers and teammate-readable packets."),
        ("Gmail", "Create drafts and schedule approved emails."),
        ("LinkedIn", "Prepare and send approved connection requests or follow-ups."),
        ("Calendar/conference sources", "Track events, speaker lists, and in-person follow-up opportunities."),
    ]
    add_table(doc, ["Connector", "Use"], connector_rows, [1.8, 4.7])

    add_h2(doc, "P4: Feedback Learning")
    add_para(
        doc,
        "Record what actually happens after outreach. The system should learn which companies, problems, people types, and message styles lead to useful conversations.",
    )
    for item in [
        "Track acceptance, reply, meeting, referral, and no-response outcomes.",
        "Compare engineer-style, recruiter-style, and FDE-style messages.",
        "Update the user's proof bank when a project framing works.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "Model and Cost Strategy")
    add_para(
        doc,
        "Use cheap models for extraction and expensive models only where judgment matters. Cost control is a product feature because the workflow can become token-heavy fast.",
    )
    cost_rows = [
        ("Cheap model tasks", "Deduping companies, extracting profile snippets, formatting source lists, word counts, first-pass summaries."),
        ("Strong model tasks", "Problem framing, technical note synthesis, claim QA, final outreach voice, and orchestration decisions."),
        ("Caching", "Cache page text, source summaries, people records, company records, and previous packet components."),
        ("Budgets", "Set per-packet token and dollar budgets. Escalate only when the packet is worth sending."),
        ("Evaluation", "Score each output for source coverage, specificity, actionability, and risk before human review."),
    ]
    add_table(doc, ["Area", "Rule"], cost_rows, [1.65, 4.85])

    add_h1(doc, "Constraints and Non-Negotiables")
    for item in [
        "No API keys or secrets in Docs, prompts, notes, screenshots, or commits. Store keys in a secret manager or local .env file. Rotate any key that has been pasted into shared text.",
        "No fabricated experience. The system can say the user is studying or prototyping a problem, but it cannot imply production experience the user does not have.",
        "No uncontrolled mass outreach. External messages require explicit approval and should be targeted.",
        "Respect platform rules and rate limits. The browser should support user-approved workflows, not scrape aggressively or bypass access controls.",
        "Every person and company claim needs a source URL or a clear uncertainty label.",
        "Every outreach draft should be under 200 words unless the user explicitly asks for a longer message.",
        "The default tone is specific, curious, student-like, and technically grounded.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "Papers and Architecture Queue")
    add_para(
        doc,
        "Paper 2405.10467 is a seed reference to review. The first task is not to force it into the product. The first task is to read it, extract the architectural relevance, and decide whether it informs the agent harness, task decomposition, evaluation, or communication protocol.",
    )
    for item in [
        "Create a paper note template: claim, mechanism, relevance, implementation idea, risk, and citation.",
        "Maintain an architecture decision record for agent communication, state, model routing, and approval gates.",
        "Prefer boring infrastructure first: task queues, structured artifacts, logs, retries, and deterministic QA.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "Teammate Execution Split")
    split_rows = [
        ("You", "Product direction, target lanes, taste, voice approval, real outreach review, and success criteria."),
        ("Teammate", "Data model, agent harness, connector plumbing, local UI or tracker, logging, and cost controls."),
        ("Codex / agent system", "Research assistance, packet drafting, document generation, verification, browser actions after approval."),
        ("Together", "Review first 20 packets, refine scoring, decide what becomes automated, and remove low-signal steps."),
    ]
    add_table(doc, ["Owner", "Responsibility"], split_rows, [1.45, 5.05])

    add_h1(doc, "First MVP Definition")
    add_para(
        doc,
        "The MVP should prove that the system can create better conversations than a normal job search. Keep the first lane narrow: AI infrastructure and inference systems.",
    )
    for item in [
        "20 sourced companies in the target lane.",
        "5 prioritized company packets.",
        "3 people per company with source-backed relevance.",
        "1 one-page technical note per company.",
        "2 outreach variants per person.",
        "Approval and send tracker.",
        "Follow-up reminders and outcome tracking.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "Immediate Next Actions")
    immediate = [
        "Freeze the target packet format and use Baseten as the gold-standard example.",
        "Create packets for Modal, Fireworks, Together, CoreWeave, and Databricks next.",
        "Build the first tracker schema with companies, people, problems, packets, messages, approvals, and follow-ups.",
        "Review paper 2405.10467 and write a one-page relevance note.",
        "Move secrets out of documents. Rotate any key that was pasted into notes or chat.",
        "Run the first 10-packet manual sprint before automating sends.",
    ]
    for item in immediate:
        add_number(doc, item)

    add_h1(doc, "Working Name")
    add_para(
        doc,
        "Opportunity Intelligence OS is the neutral working name. Better product names can come later. The important thing now is the operating loop: source, understand, map proof, write, verify, approve, send, and learn.",
    )

    doc.save(DOCX)
    doc.save(GDOCS_DOCX)


if __name__ == "__main__":
    build()
    print(DOCX)
    print(GDOCS_DOCX)
